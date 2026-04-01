#!/usr/bin/env python3
"""
Generate WMC Anforderungsportal User Manuals (EN + DE) as .docx files.
Run: "C:/Program Files/Python311/python.exe" scripts/generate-user-manuals.py
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "docs")
os.makedirs(OUTPUT_DIR, exist_ok=True)


# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set background shading of a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_caption(doc, caption_text, table_number):
    """Add a table caption below/above a table."""
    p = doc.add_paragraph()
    p.style = doc.styles["Caption"] if "Caption" in [s.name for s in doc.styles] else doc.styles["Normal"]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"Table {table_number}: {caption_text}")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    return p


def create_styled_table(doc, headers, rows, table_num, caption):
    """Create a professional table with header styling and caption."""
    add_table_caption(doc, caption, table_num)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1A365D")
    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph("")  # spacing
    return table


def add_toc_field(doc):
    """Insert a Word TOC field that updates on open."""
    p = doc.add_paragraph()
    run = p.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar1)
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instrText)
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    run._r.append(fldChar2)
    run2 = p.add_run("(Right-click and select 'Update Field' to refresh the Table of Contents)")
    run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run2.font.size = Pt(10)
    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run2._r.append(fldChar3)


def setup_document(doc, lang):
    """Configure document styles and defaults."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level in range(1, 4):
        h_style = doc.styles[f"Heading {level}"]
        h_style.font.name = "Calibri"
        h_style.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        if level == 1:
            h_style.font.size = Pt(20)
        elif level == 2:
            h_style.font.size = Pt(16)
        else:
            h_style.font.size = Pt(13)

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)


def add_cover_page(doc, lang):
    """Create a professional cover page."""
    for _ in range(4):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("WMC Anforderungsportal")
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if lang == "en":
        run = subtitle.add_run("User Manual")
    else:
        run = subtitle.add_run("Benutzerhandbuch")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x4A, 0x5A, 0x7A)

    doc.add_paragraph("")

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if lang == "en":
        desc_text = "Complete guide for using the WMC Requirements Portal\nfor project requirement collection, AI-powered forms, and administration."
    else:
        desc_text = "Vollständige Anleitung zur Nutzung des WMC-Anforderungsportals\nfür die Erfassung von Projektanforderungen, KI-gestützte Formulare und Administration."
    run = desc.add_run(desc_text)
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for _ in range(6):
        doc.add_paragraph("")

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lines = [
        ("Version:", "1.0"),
        ("Date:" if lang == "en" else "Datum:", "March 2026" if lang == "en" else "März 2026"),
        ("Author:" if lang == "en" else "Autor:", "WMC — Wamocon Consulting"),
        ("Classification:" if lang == "en" else "Klassifikation:", "Internal / Client Use" if lang == "en" else "Intern / Kundennutzung"),
    ]
    for label, value in lines:
        run = meta.add_run(f"{label} {value}\n")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()


# ─────────────────────────────────────────────────────────────────────────────
# Content: ENGLISH
# ─────────────────────────────────────────────────────────────────────────────

def build_english(doc):
    table_num = [0]  # mutable counter

    def tn():
        table_num[0] += 1
        return table_num[0]

    setup_document(doc, "en")
    add_cover_page(doc, "en")

    # ── TABLE OF CONTENTS ────────────────────────────────────────────────────
    doc.add_heading("Table of Contents", level=1)
    add_toc_field(doc)
    doc.add_page_break()

    # ── LIST OF TABLES ───────────────────────────────────────────────────────
    doc.add_heading("List of Tables", level=1)
    lot_items = [
        "Table 1: User Roles Overview",
        "Table 2: Response Status Values",
        "Table 3: Supported Languages",
        "Table 4: Question Types in Templates",
        "Table 5: Navigation Overview — Admin",
        "Table 6: Navigation Overview — Client",
        "Table 7: System Requirements",
        "Table 8: Keyboard Shortcuts",
        "Table 9: Troubleshooting Guide",
    ]
    for item in lot_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # CHAPTER 1: INTRODUCTION
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("1  Introduction", level=1)

    doc.add_heading("1.1  What Is the WMC Anforderungsportal?", level=2)
    doc.add_paragraph(
        "The WMC Anforderungsportal (Requirements Portal) is a modern web application "
        "developed by WMC — Wamocon Consulting. It helps businesses and consultants "
        "collect, organise, and review project requirements in a structured and efficient way."
    )
    doc.add_paragraph(
        "Instead of exchanging long email chains, spreadsheets, or unstructured documents, "
        "the portal provides smart digital forms that guide idea owners step-by-step through "
        "defining their project vision, target audience, features, and budget."
    )

    doc.add_heading("1.2  What Problem Does It Solve?", level=2)
    doc.add_paragraph(
        "Collecting requirements for software projects is traditionally chaotic. "
        "Clients often struggle to express their ideas clearly, information gets lost "
        "in emails, and consultants spend weeks piecing together incomplete specifications."
    )
    doc.add_paragraph("The Anforderungsportal solves this by providing:")
    items = [
        "Structured, template-based forms that ensure nothing is missed",
        "AI-powered text polishing that turns rough notes into professional requirement texts",
        "Voice input so you can simply speak your ideas instead of typing",
        "An AI Interviewer that asks smart follow-up questions",
        "Multi-language support (25 languages) for international teams",
        "A review dashboard for consultants to track and evaluate submissions",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("1.3  Who Is This Manual For?", level=2)
    doc.add_paragraph(
        "This manual is intended for two main groups of users:"
    )
    doc.add_paragraph(
        "Product Owners / Idea Owners (Clients): People who have an app or software idea and "
        "want to submit their requirements through the portal.", style="List Bullet"
    )
    doc.add_paragraph(
        "WMC Staff / Administrators: Consultants and managers who create projects, "
        "design templates, and review submitted requirements.", style="List Bullet"
    )

    doc.add_heading("1.4  System Requirements", level=2)
    create_styled_table(doc,
        ["Requirement", "Details"],
        [
            ["Web Browser", "Google Chrome, Mozilla Firefox, Microsoft Edge, or Safari (latest version)"],
            ["Internet Connection", "Stable broadband connection required"],
            ["Screen Size", "Optimised for desktop; works on tablets and smartphones"],
            ["Microphone", "Required for voice input feature (optional)"],
            ["Operating System", "Any (Windows, macOS, Linux, iOS, Android)"],
        ],
        tn(), "System Requirements"
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # CHAPTER 2: GETTING STARTED
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("2  Getting Started", level=1)

    doc.add_heading("2.1  Accessing the Portal", level=2)
    doc.add_paragraph(
        "Open your web browser and navigate to the portal URL provided by your WMC contact. "
        "The address typically looks like: https://anforderungsportal.wamocon.com"
    )
    doc.add_paragraph(
        "You will see the landing page with an overview of the portal's features, "
        "trust badges (GDPR Ready, Made in Germany), and buttons to get started."
    )

    doc.add_heading("2.2  Creating an Account / Logging In", level=2)
    doc.add_paragraph(
        "The portal supports two ways to log in:"
    )
    doc.add_heading("2.2.1  Magic Link (Recommended)", level=3)
    doc.add_paragraph(
        "1. Click 'Get Started' or 'Login' on the landing page.\n"
        "2. Enter your email address in the login form.\n"
        "3. Click 'Send Magic Link'.\n"
        "4. Check your email inbox for a message from the portal.\n"
        "5. Click the link in the email — you will be automatically logged in.\n"
        "6. The magic link is valid for a limited time. If expired, simply request a new one."
    )

    doc.add_heading("2.2.2  Password Login", level=3)
    doc.add_paragraph(
        "If you have been given a password by your administrator:\n"
        "1. Click 'Login' on the landing page.\n"
        "2. Enter your email and password.\n"
        "3. Click the 'Sign In' button.\n"
        "4. You will be redirected to your personal dashboard."
    )

    doc.add_heading("2.2.3  Forgot Password", level=3)
    doc.add_paragraph(
        "1. On the login page, click 'Forgot Password?'\n"
        "2. Enter your registered email address.\n"
        "3. You will receive an email with a link to reset your password.\n"
        "4. Follow the link and create a new password."
    )

    doc.add_heading("2.3  Understanding User Roles", level=2)
    doc.add_paragraph(
        "The portal has different roles that determine what you can see and do:"
    )
    create_styled_table(doc,
        ["Role", "What You Can Do", "Where You Land After Login"],
        [
            ["Super Admin", "Full access to all settings, projects, templates, and users", "Admin Dashboard"],
            ["Staff (WMC Consultant)", "Create and manage projects, templates; review client submissions", "Admin Dashboard"],
            ["Product Owner (Client)", "Fill out requirement forms, use AI interviewer, submit responses", "My Projects"],
        ],
        tn(), "User Roles Overview"
    )

    doc.add_heading("2.4  Choosing Your Language", level=2)
    doc.add_paragraph(
        "The portal supports 25 languages. To change the language:\n"
        "1. Look for the language selector in the top navigation bar (it shows a flag or language code).\n"
        "2. Click on it to open the language dropdown.\n"
        "3. Select your preferred language.\n"
        "4. The entire interface will switch to your chosen language immediately."
    )

    doc.add_heading("2.5  Dark Mode / Light Mode", level=2)
    doc.add_paragraph(
        "You can switch between a light and dark visual theme:\n"
        "1. Find the sun/moon icon in the top navigation bar.\n"
        "2. Click it to toggle between light mode (white background) and dark mode (dark background).\n"
        "3. Your preference is saved automatically."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # CHAPTER 3: CLIENT / PRODUCT OWNER GUIDE
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("3  Guide for Product Owners (Clients)", level=1)

    doc.add_heading("3.1  My Projects Dashboard", level=2)
    doc.add_paragraph(
        "After logging in as a product owner, you will see the 'My Projects' page. "
        "This page lists all projects you have been invited to or assigned to."
    )
    doc.add_paragraph(
        "Each project card shows:\n"
        "• The project name\n"
        "• The current status (e.g., 'In Progress', 'Submitted')\n"
        "• A progress indicator showing how many questions you have answered\n"
        "• A button to continue filling out the form or review your submission"
    )

    doc.add_heading("3.2  Filling Out a Requirement Form", level=2)
    doc.add_paragraph(
        "Click on a project to open the requirement form. The form is divided into "
        "logical sections (e.g., General Idea, Target Group, Technical Features, Budget)."
    )

    doc.add_heading("3.2.1  Step-by-Step Form Navigation", level=3)
    doc.add_paragraph(
        "1. The form guides you section by section.\n"
        "2. Fill in each field — text, selections, dates, or file uploads.\n"
        "3. Click 'Next' to move to the next section or 'Back' to return to a previous one.\n"
        "4. Your answers are saved automatically as you type.\n"
        "5. You can leave the form at any time and return later — your progress is preserved."
    )

    doc.add_heading("3.2.2  Using Voice Input", level=3)
    doc.add_paragraph(
        "Instead of typing, you can speak your answers:"
    )
    doc.add_paragraph(
        "1. Look for the microphone icon next to text fields.\n"
        "2. Click the microphone icon.\n"
        "3. Allow your browser to access your microphone (first time only).\n"
        "4. Speak clearly — your words will appear as text in real time.\n"
        "5. Click the stop button when finished.\n"
        "6. The system will automatically clean up and polish your spoken text."
    )
    doc.add_paragraph(
        "Tip: Voice input works best in a quiet environment. Speak in complete sentences for the best results."
    )

    doc.add_heading("3.2.3  AI Text Polishing", level=3)
    doc.add_paragraph(
        "The portal includes an intelligent text improvement feature powered by AI (Google Gemini). "
        "Whenever you enter text (either by typing or voice), the system can:"
    )
    doc.add_paragraph("Correct grammar and spelling mistakes", style="List Bullet")
    doc.add_paragraph("Improve the professional tone of your text", style="List Bullet")
    doc.add_paragraph("Structure your thoughts into clear requirement statements", style="List Bullet")
    doc.add_paragraph(
        "The polishing happens automatically or can be triggered manually. "
        "You always see both your original text and the polished version, "
        "so you can approve or discard the changes."
    )

    doc.add_heading("3.2.4  AI Follow-Up Questions", level=3)
    doc.add_paragraph(
        "After you answer a question, the AI may suggest follow-up questions to help you "
        "provide more detail. These appear as highlighted cards below your answer."
    )
    doc.add_paragraph(
        "You can choose to answer these follow-ups or skip them — they are optional "
        "but help create better, more complete requirements."
    )

    doc.add_heading("3.2.5  Uploading Files", level=3)
    doc.add_paragraph(
        "Some questions allow or require file uploads (e.g., mockups, wireframes, or reference documents).\n"
        "1. Click the upload area or drag and drop a file.\n"
        "2. Supported formats include PDF, images (PNG, JPG), and common office documents.\n"
        "3. Maximum file size: 50 MB per file.\n"
        "4. Uploaded files appear as thumbnail previews below the question."
    )

    doc.add_heading("3.3  Using the AI Interviewer", level=2)
    doc.add_paragraph(
        "The AI Interviewer is an alternative way to provide your requirements through a "
        "natural conversation instead of filling form fields."
    )
    doc.add_paragraph(
        "1. Open your project form.\n"
        "2. Look for the 'AI Interview' or 'Chat' mode button.\n"
        "3. Switch to the interview mode.\n"
        "4. A friendly AI bot will ask you questions one by one.\n"
        "5. You can type your answers or use voice input.\n"
        "6. The AI understands context and asks relevant follow-up questions.\n"
        "7. When the interview is complete, all your answers are automatically mapped "
        "back to the form fields."
    )
    doc.add_paragraph(
        "Tip: The AI Interviewer is ideal if you prefer a conversational approach "
        "or find it difficult to fill out structured forms."
    )

    doc.add_heading("3.4  Reviewing and Submitting", level=2)
    doc.add_paragraph(
        "Before final submission:\n"
        "1. Navigate to the 'Review' section of the form.\n"
        "2. You will see a full summary of all your answers, section by section.\n"
        "3. Check everything carefully — once submitted, you cannot easily change answers.\n"
        "4. If you want to change something, click on the section name to go back and edit.\n"
        "5. When you are satisfied, click 'Submit'.\n"
        "6. You will see a confirmation message. Your submission is now visible to the WMC team."
    )

    doc.add_heading("3.5  After Submission", level=2)
    doc.add_paragraph(
        "After submitting, your response status changes to 'Submitted'. "
        "The WMC team will review your requirements and may:\n"
        "• Mark the response as 'Reviewed' (no further action needed from you)\n"
        "• Request additional information (you will be notified)\n\n"
        "You can check the status of your submission at any time in the 'My Projects' dashboard."
    )

    create_styled_table(doc,
        ["Status", "Meaning"],
        [
            ["Draft", "You have started filling the form but have not submitted yet"],
            ["In Progress", "You are actively working on the form"],
            ["Submitted", "You have submitted; the WMC team is reviewing"],
            ["Reviewed", "The WMC team has reviewed your submission"],
        ],
        tn(), "Response Status Values"
    )

    doc.add_heading("3.6  Account Settings", level=2)
    doc.add_paragraph(
        "Access your account settings by clicking your profile icon in the top-right corner "
        "or navigating to the 'Account' page."
    )
    doc.add_paragraph(
        "Here you can:\n"
        "• See your assigned role and account information\n"
        "• Change your password\n"
        "• View how long you have been a member"
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # CHAPTER 4: ADMIN / STAFF GUIDE
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("4  Guide for Administrators (WMC Staff)", level=1)

    doc.add_heading("4.1  Admin Dashboard Overview", level=2)
    doc.add_paragraph(
        "After logging in as a staff member or super admin, you land on the Admin Dashboard. "
        "It provides a quick overview of your organisation's activity."
    )
    doc.add_paragraph("The dashboard shows the following key numbers:")
    doc.add_paragraph("Total Projects — number of all projects in your organisation", style="List Bullet")
    doc.add_paragraph("Active Projects — projects currently in progress", style="List Bullet")
    doc.add_paragraph("Total Responses — all submissions received from clients", style="List Bullet")
    doc.add_paragraph("Pending Review — submissions waiting for your review (highlighted in red)", style="List Bullet")
    doc.add_paragraph(
        "Below the numbers, you can see the most recent activity — latest submissions, "
        "status changes, and quick-access buttons to common tasks."
    )

    create_styled_table(doc,
        ["Menu Item", "Description"],
        [
            ["Dashboard", "Overview with key metrics and recent activity"],
            ["Projects", "Create, manage, and assign requirement collection projects"],
            ["Templates", "Design question templates used in project forms"],
            ["Responses", "Review, filter, and export client submissions"],
            ["Archive", "View completed or inactive projects"],
            ["Settings", "Organisation and system settings"],
        ],
        tn(), "Navigation Overview — Admin"
    )

    doc.add_heading("4.2  Managing Projects", level=2)

    doc.add_heading("4.2.1  Creating a New Project", level=3)
    doc.add_paragraph(
        "1. Go to 'Projects' in the sidebar menu.\n"
        "2. Click the 'New Project' button.\n"
        "3. Fill in the project details:\n"
        "   • Project Name: A clear, descriptive name (e.g., 'Travel App Requirements')\n"
        "   • Slug: An automatically generated URL-friendly identifier\n"
        "   • Deadline: Set the number of days for completion\n"
        "   • Template: Choose which question template to use\n"
        "   • Welcome Text: A greeting message clients see when they open the form\n"
        "4. Click 'Create Project'.\n"
        "5. The project is now ready to receive submissions."
    )

    doc.add_heading("4.2.2  Inviting Team Members", level=3)
    doc.add_paragraph(
        "After creating a project, you can invite clients (product owners) to fill it out:\n"
        "1. Open the project detail page.\n"
        "2. Click 'Invite Member' or 'Add Member'.\n"
        "3. Enter the email address of the person you want to invite.\n"
        "4. Choose their role (Product Owner or Client).\n"
        "5. A magic link invitation email will be sent automatically.\n"
        "6. The invitation is valid for 7 days."
    )

    doc.add_heading("4.2.3  Approving Project Proposals", level=3)
    doc.add_paragraph(
        "When clients submit project ideas, they may require approval:\n"
        "1. Open the project detail page.\n"
        "2. Review the project description and proposed scope.\n"
        "3. Click 'Approve' to activate the project, or\n"
        "4. Click 'Request Changes' to send feedback with a staff note."
    )

    doc.add_heading("4.3  Managing Templates", level=2)
    doc.add_paragraph(
        "Templates define the structure of the requirement form that clients fill out. "
        "A template consists of sections, and each section contains questions."
    )

    doc.add_heading("4.3.1  Creating a Template", level=3)
    doc.add_paragraph(
        "1. Go to 'Templates' in the sidebar.\n"
        "2. Click 'New Template'.\n"
        "3. Give it a name and optional description.\n"
        "4. Start adding sections (e.g., 'General Information', 'Technical Requirements').\n"
        "5. Within each section, add questions of various types."
    )

    create_styled_table(doc,
        ["Question Type", "Description", "Example Use"],
        [
            ["Short Text", "Single-line text input", "Project name, company name"],
            ["Long Text", "Multi-line text area", "Detailed descriptions, ideas"],
            ["Single Choice", "Radio buttons — pick one option", "Priority level (High/Medium/Low)"],
            ["Multiple Choice", "Checkboxes — pick several options", "Desired platforms (Web, iOS, Android)"],
            ["Dropdown", "Select from a list", "Industry sector"],
            ["Date", "Date picker", "Desired launch date"],
            ["File Upload", "Attach documents or images", "Mockups, wireframes, reference PDFs"],
        ],
        tn(), "Question Types in Templates"
    )

    doc.add_heading("4.3.2  Reordering Sections and Questions", level=3)
    doc.add_paragraph(
        "You can change the order of sections and questions using drag-and-drop:\n"
        "1. Hover over the drag handle (≡ icon) next to a section or question.\n"
        "2. Click and hold, then drag it to the desired position.\n"
        "3. Release to drop it in its new location.\n"
        "4. Changes are saved automatically."
    )

    doc.add_heading("4.4  Reviewing Responses", level=2)
    doc.add_paragraph(
        "When clients submit their requirements, you can review them in the Responses area."
    )
    doc.add_paragraph(
        "1. Go to 'Responses' in the sidebar.\n"
        "2. You will see a list of all submissions with their status and progress.\n"
        "3. Use the filters to find specific responses:\n"
        "   • Filter by status (Draft, In Progress, Submitted, Reviewed)\n"
        "   • Search by respondent name or project name\n"
        "4. Click on a response to view the full details.\n"
        "5. Review each answer carefully.\n"
        "6. Mark the response as 'Reviewed' when you are done."
    )

    doc.add_heading("4.4.1  Exporting Responses", level=3)
    doc.add_paragraph(
        "You can export response data for offline processing:\n"
        "1. In the Responses list, click the 'Export CSV' button.\n"
        "2. A CSV file will be downloaded containing all response data.\n"
        "3. Open it in Microsoft Excel, Google Sheets, or any spreadsheet application."
    )

    doc.add_heading("4.5  Archive", level=2)
    doc.add_paragraph(
        "The Archive section stores completed or deactivated projects. "
        "Projects are moved here when they are no longer active. "
        "You can still view all details and responses of archived projects."
    )

    doc.add_heading("4.6  Settings", level=2)
    doc.add_paragraph(
        "The Settings page allows administrators to configure organisation-level settings. "
        "Available options may include organisation name, default language, and notification preferences."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # CHAPTER 5: NAVIGATION REFERENCE
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("5  Navigation Reference", level=1)

    doc.add_heading("5.1  Client Navigation", level=2)
    create_styled_table(doc,
        ["Page", "How to Access", "What You Can Do"],
        [
            ["My Projects", "Automatic after login", "See all your assigned projects"],
            ["Form (Fill Mode)", "Click a project", "Answer questions section by section"],
            ["Form (Interview Mode)", "Switch mode in form", "Answer questions through AI conversation"],
            ["Review", "Last step in form", "Review all answers before submitting"],
            ["Account", "Click profile icon", "View account info, change password"],
        ],
        tn(), "Navigation Overview — Client"
    )

    doc.add_heading("5.2  Keyboard Shortcuts", level=2)
    create_styled_table(doc,
        ["Shortcut", "Action"],
        [
            ["Tab", "Move to the next form field"],
            ["Shift + Tab", "Move to the previous form field"],
            ["Enter", "Confirm selection or submit current field"],
            ["Ctrl/Cmd + S", "Save current progress (in supported browsers)"],
            ["Esc", "Close dialogs or pop-ups"],
        ],
        tn(), "Keyboard Shortcuts"
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # CHAPTER 6: LANGUAGES
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("6  Supported Languages", level=1)
    doc.add_paragraph(
        "The WMC Anforderungsportal supports 25 languages. The default language is German (DE). "
        "All interface elements, form labels, and AI-generated text can appear in the selected language."
    )
    create_styled_table(doc,
        ["Code", "Language", "Code", "Language"],
        [
            ["de", "German (Deutsch)", "nl", "Dutch (Nederlands)"],
            ["en", "English", "no", "Norwegian (Norsk)"],
            ["bg", "Bulgarian (Български)", "pl", "Polish (Polski)"],
            ["cs", "Czech (Čeština)", "pt", "Portuguese (Português)"],
            ["da", "Danish (Dansk)", "ro", "Romanian (Română)"],
            ["el", "Greek (Ελληνικά)", "ru", "Russian (Русский)"],
            ["es", "Spanish (Español)", "sk", "Slovak (Slovenčina)"],
            ["et", "Estonian (Eesti)", "sl", "Slovenian (Slovenščina)"],
            ["fi", "Finnish (Suomi)", "sv", "Swedish (Svenska)"],
            ["fr", "French (Français)", "tr", "Turkish (Türkçe)"],
            ["hr", "Croatian (Hrvatski)", "hu", "Hungarian (Magyar)"],
            ["it", "Italian (Italiano)", "lt", "Lithuanian (Lietuvių)"],
            ["lv", "Latvian (Latviešu)", "", ""],
        ],
        tn(), "Supported Languages"
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # CHAPTER 7: DATA PRIVACY & SECURITY
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("7  Data Privacy and Security", level=1)

    doc.add_heading("7.1  GDPR Compliance", level=2)
    doc.add_paragraph(
        "The WMC Anforderungsportal is fully compliant with the European General Data Protection "
        "Regulation (GDPR / DSGVO). Key measures include:"
    )
    doc.add_paragraph("All data is stored on servers within the European Union", style="List Bullet")
    doc.add_paragraph("Data transmission is encrypted using HTTPS/TLS", style="List Bullet")
    doc.add_paragraph("Access is controlled through role-based permissions", style="List Bullet")
    doc.add_paragraph("Users can request deletion of their data at any time", style="List Bullet")
    doc.add_paragraph("AI processing is done securely — no data is used for model training", style="List Bullet")

    doc.add_heading("7.2  Data Storage", level=2)
    doc.add_paragraph(
        "The portal uses Supabase as its backend database service. "
        "Supabase provides enterprise-grade PostgreSQL databases with built-in "
        "authentication, row-level security, and automatic backups."
    )

    doc.add_heading("7.3  Authentication Security", level=2)
    doc.add_paragraph(
        "Login is protected through:\n"
        "• Multi-factor options (magic links, passwords)\n"
        "• Encrypted password storage (bcrypt hashing)\n"
        "• Token-based session management (JWT)\n"
        "• Automatic session expiry after periods of inactivity"
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # CHAPTER 8: TROUBLESHOOTING
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("8  Troubleshooting", level=1)
    doc.add_paragraph(
        "If you encounter problems while using the portal, try the solutions below "
        "before contacting support."
    )

    create_styled_table(doc,
        ["Problem", "Possible Cause", "Solution"],
        [
            [
                "Cannot log in",
                "Wrong email or password",
                "Double-check your credentials. Use 'Forgot Password' to reset. "
                "If using a magic link, check your spam/junk folder."
            ],
            [
                "Page is not loading",
                "Internet connection issue or browser cache",
                "Check your internet connection. Try clearing your browser cache "
                "(Ctrl+Shift+Delete) or use a private/incognito window."
            ],
            [
                "Voice input is not working",
                "Microphone not permitted",
                "Your browser needs permission to access your microphone. "
                "Click the lock icon in the address bar and allow microphone access. "
                "Make sure no other application is using the microphone."
            ],
            [
                "AI Polish is not improving my text",
                "Text too short or unclear",
                "Try providing at least 2-3 complete sentences. The AI works best "
                "with longer, more detailed inputs."
            ],
            [
                "File upload fails",
                "File too large or unsupported format",
                "Maximum file size is 50 MB. Supported formats: PDF, PNG, JPG, "
                "and common office documents."
            ],
            [
                "My progress was not saved",
                "Session expired",
                "The system auto-saves, but extended periods of inactivity may "
                "log you out. Log back in — most of your progress should be preserved."
            ],
            [
                "Language changed unexpectedly",
                "Browser language detection",
                "The portal auto-detects your browser language. Use the language "
                "switcher in the top bar to manually select your preferred language."
            ],
            [
                "Cannot see my project",
                "Not assigned to the project",
                "Contact your WMC administrator to confirm you have been invited "
                "to the correct project."
            ],
        ],
        tn(), "Troubleshooting Guide"
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # CHAPTER 9: FREQUENTLY ASKED QUESTIONS
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("9  Frequently Asked Questions (FAQ)", level=1)

    faqs = [
        ("Can I edit my responses after submitting?",
         "Once submitted, your responses are locked for review by the WMC team. "
         "If you need to make changes, contact your WMC project manager, who can "
         "reopen the form for you."),
        ("Is my data safe?",
         "Yes. All data is encrypted in transit and at rest. The portal complies with "
         "GDPR, and your data is stored on EU-based servers. AI features do not use "
         "your data for training purposes."),
        ("Can I use the portal on my phone?",
         "Yes. The portal is responsive and works on smartphones and tablets. "
         "For the best experience, we recommend using a desktop or laptop computer, "
         "especially for longer forms."),
        ("What happens if my internet disconnects while filling the form?",
         "Your answers are saved automatically as you type. If you lose connection, "
         "simply reconnect and reload the page — your progress should still be there."),
        ("Do I need to install any software?",
         "No. The portal runs entirely in your web browser. No downloads or installations "
         "are required."),
        ("How many languages are supported?",
         "The portal supports 25 languages, including German, English, French, Spanish, "
         "Italian, Polish, Turkish, and many more. See Chapter 6 for the full list."),
        ("Who can see my submitted requirements?",
         "Only the WMC staff members and administrators assigned to your project can see "
         "your submissions. Other clients cannot see your data."),
        ("Can I use voice input in any language?",
         "Voice input depends on your browser's speech recognition capabilities. "
         "Most modern browsers support major European languages. The best results "
         "are achieved with Chrome."),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        run = p.add_run(f"Q: {q}")
        run.bold = True
        doc.add_paragraph(f"A: {a}")
        doc.add_paragraph("")  # spacing

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # CHAPTER 10: GLOSSARY
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("10  Glossary", level=1)
    terms = [
        ("Anforderungsportal", "The German name for 'Requirements Portal' — the application described in this manual."),
        ("AI Interviewer", "A chat-based assistant that asks you questions to help collect requirements."),
        ("AI Polish", "An automatic feature that improves the grammar and tone of your text."),
        ("Dashboard", "The overview page showing key numbers and recent activity."),
        ("GDPR / DSGVO", "European data protection regulation ensuring privacy of personal data."),
        ("Magic Link", "A secure, one-time login link sent to your email — no password needed."),
        ("Product Owner", "The person who has the product idea and provides requirements."),
        ("Response", "The set of answers a client submits for a project."),
        ("Section", "A group of related questions within a template."),
        ("Slug", "A URL-friendly version of a project name (e.g., 'travel-app-requirements')."),
        ("Staff", "WMC consultants who manage projects and review submissions."),
        ("Super Admin", "A system administrator with full access to all features."),
        ("Template", "A pre-defined structure of sections and questions used for collecting requirements."),
        ("Voice Input", "Speaking instead of typing — the portal converts your speech to text."),
    ]
    for term, defn in terms:
        p = doc.add_paragraph()
        run = p.add_run(f"{term}: ")
        run.bold = True
        p.add_run(defn)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # CHAPTER 11: SUPPORT & CONTACT
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("11  Support and Contact", level=1)
    doc.add_paragraph(
        "If you need further assistance, please contact the WMC support team:"
    )
    doc.add_paragraph("Email: support@wamocon.com", style="List Bullet")
    doc.add_paragraph("Website: https://wamocon.com", style="List Bullet")
    doc.add_paragraph("Phone: Contact your project manager for direct assistance", style="List Bullet")
    doc.add_paragraph("")
    doc.add_paragraph(
        "When contacting support, please include:\n"
        "• Your email address (used for login)\n"
        "• The project name\n"
        "• A description of the issue\n"
        "• Screenshots if available"
    )

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— End of User Manual —")
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.size = Pt(10)


# ─────────────────────────────────────────────────────────────────────────────
# Content: GERMAN
# ─────────────────────────────────────────────────────────────────────────────

def build_german(doc):
    table_num = [0]

    def tn():
        table_num[0] += 1
        return table_num[0]

    setup_document(doc, "de")
    add_cover_page(doc, "de")

    # ── INHALTSVERZEICHNIS ────────────────────────────────────────────────
    doc.add_heading("Inhaltsverzeichnis", level=1)
    add_toc_field(doc)
    doc.add_page_break()

    # ── TABELLENVERZEICHNIS ──────────────────────────────────────────────
    doc.add_heading("Tabellenverzeichnis", level=1)
    lot_items = [
        "Tabelle 1: Übersicht der Benutzerrollen",
        "Tabelle 2: Status-Werte der Einreichungen",
        "Tabelle 3: Unterstützte Sprachen",
        "Tabelle 4: Fragetypen in Vorlagen",
        "Tabelle 5: Navigationsübersicht — Administrator",
        "Tabelle 6: Navigationsübersicht — Auftraggeber",
        "Tabelle 7: Systemanforderungen",
        "Tabelle 8: Tastenkombinationen",
        "Tabelle 9: Fehlerbehebung",
    ]
    for item in lot_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # KAPITEL 1: EINFÜHRUNG
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("1  Einführung", level=1)

    doc.add_heading("1.1  Was ist das WMC-Anforderungsportal?", level=2)
    doc.add_paragraph(
        "Das WMC-Anforderungsportal ist eine moderne Webanwendung, entwickelt von "
        "WMC — Wamocon Consulting. Es unterstützt Unternehmen und Berater dabei, "
        "Projektanforderungen strukturiert, effizient und vollständig zu erfassen."
    )
    doc.add_paragraph(
        "Statt langer E-Mail-Ketten, Tabellenkalkulationen oder unstrukturierter Dokumente "
        "bietet das Portal intelligente digitale Formulare, die Ideengeber Schritt für Schritt "
        "durch die Definition ihrer Projektvision, Zielgruppe, Funktionen und Budget führen."
    )

    doc.add_heading("1.2  Welches Problem wird gelöst?", level=2)
    doc.add_paragraph(
        "Die Erfassung von Anforderungen für Softwareprojekte ist traditionell chaotisch. "
        "Kunden haben oft Schwierigkeiten, ihre Ideen klar auszudrücken, Informationen gehen "
        "in E-Mails verloren, und Berater verbringen Wochen damit, unvollständige Spezifikationen "
        "zusammenzusetzen."
    )
    doc.add_paragraph("Das Anforderungsportal löst dies durch:")
    items = [
        "Strukturierte, vorlagenbasierte Formulare, die sicherstellen, dass nichts vergessen wird",
        "KI-gestützte Textverbesserung, die grobe Notizen in professionelle Anforderungstexte verwandelt",
        "Spracheingabe, mit der Sie Ihre Ideen einfach sprechen können, statt zu tippen",
        "Einen KI-Interviewer, der intelligente Nachfragen stellt",
        "Mehrsprachige Unterstützung (25 Sprachen) für internationale Teams",
        "Ein Prüfungs-Dashboard für Berater zur Nachverfolgung und Bewertung von Einreichungen",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("1.3  Für wen ist dieses Handbuch?", level=2)
    doc.add_paragraph("Dieses Handbuch richtet sich an zwei Hauptgruppen von Benutzern:")
    doc.add_paragraph(
        "Product Owner / Ideengeber (Auftraggeber): Personen, die eine App- oder Software-Idee haben "
        "und ihre Anforderungen über das Portal einreichen möchten.", style="List Bullet"
    )
    doc.add_paragraph(
        "WMC-Mitarbeiter / Administratoren: Berater und Manager, die Projekte erstellen, "
        "Vorlagen gestalten und eingereichte Anforderungen prüfen.", style="List Bullet"
    )

    doc.add_heading("1.4  Systemanforderungen", level=2)
    create_styled_table(doc,
        ["Anforderung", "Details"],
        [
            ["Webbrowser", "Google Chrome, Mozilla Firefox, Microsoft Edge oder Safari (aktuelle Version)"],
            ["Internetverbindung", "Stabile Breitbandverbindung erforderlich"],
            ["Bildschirmgröße", "Optimiert für Desktop; funktioniert auf Tablets und Smartphones"],
            ["Mikrofon", "Erforderlich für die Spracheingabe-Funktion (optional)"],
            ["Betriebssystem", "Beliebig (Windows, macOS, Linux, iOS, Android)"],
        ],
        tn(), "Systemanforderungen"
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # KAPITEL 2: ERSTE SCHRITTE
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("2  Erste Schritte", level=1)

    doc.add_heading("2.1  Zugang zum Portal", level=2)
    doc.add_paragraph(
        "Öffnen Sie Ihren Webbrowser und navigieren Sie zur Portal-URL, die Sie von "
        "Ihrem WMC-Ansprechpartner erhalten haben. Die Adresse lautet typischerweise: "
        "https://anforderungsportal.wamocon.com"
    )
    doc.add_paragraph(
        "Sie sehen die Startseite mit einer Übersicht der Portal-Funktionen, "
        "Vertrauenssiegeln (DSGVO-konform, Made in Germany) und Schaltflächen zum Starten."
    )

    doc.add_heading("2.2  Konto erstellen / Anmelden", level=2)
    doc.add_paragraph("Das Portal unterstützt zwei Anmeldemethoden:")

    doc.add_heading("2.2.1  Magic Link (Empfohlen)", level=3)
    doc.add_paragraph(
        "1. Klicken Sie auf 'Jetzt starten' oder 'Anmelden' auf der Startseite.\n"
        "2. Geben Sie Ihre E-Mail-Adresse ein.\n"
        "3. Klicken Sie auf 'Magic Link senden'.\n"
        "4. Prüfen Sie Ihren E-Mail-Posteingang auf eine Nachricht vom Portal.\n"
        "5. Klicken Sie auf den Link in der E-Mail — Sie werden automatisch angemeldet.\n"
        "6. Der Magic Link ist zeitlich begrenzt gültig. Falls abgelaufen, fordern Sie einfach einen neuen an."
    )

    doc.add_heading("2.2.2  Passwort-Anmeldung", level=3)
    doc.add_paragraph(
        "Wenn Sie ein Passwort von Ihrem Administrator erhalten haben:\n"
        "1. Klicken Sie auf 'Anmelden' auf der Startseite.\n"
        "2. Geben Sie Ihre E-Mail-Adresse und Ihr Passwort ein.\n"
        "3. Klicken Sie auf 'Anmelden'.\n"
        "4. Sie werden zu Ihrem persönlichen Dashboard weitergeleitet."
    )

    doc.add_heading("2.2.3  Passwort vergessen", level=3)
    doc.add_paragraph(
        "1. Klicken Sie auf der Anmeldeseite auf 'Passwort vergessen?'\n"
        "2. Geben Sie Ihre registrierte E-Mail-Adresse ein.\n"
        "3. Sie erhalten eine E-Mail mit einem Link zum Zurücksetzen.\n"
        "4. Folgen Sie dem Link und erstellen Sie ein neues Passwort."
    )

    doc.add_heading("2.3  Benutzerrollen verstehen", level=2)
    doc.add_paragraph("Das Portal verfügt über verschiedene Rollen, die festlegen, was Sie sehen und tun können:")
    create_styled_table(doc,
        ["Rolle", "Was Sie tun können", "Wo Sie nach der Anmeldung landen"],
        [
            ["Super-Admin", "Vollzugriff auf alle Einstellungen, Projekte, Vorlagen und Benutzer", "Admin-Dashboard"],
            ["Mitarbeiter (WMC-Berater)", "Projekte und Vorlagen erstellen und verwalten; Kundeneinreichungen prüfen", "Admin-Dashboard"],
            ["Product Owner (Auftraggeber)", "Anforderungsformulare ausfüllen, KI-Interviewer nutzen, Antworten einreichen", "Meine Projekte"],
        ],
        tn(), "Übersicht der Benutzerrollen"
    )

    doc.add_heading("2.4  Sprache auswählen", level=2)
    doc.add_paragraph(
        "Das Portal unterstützt 25 Sprachen. Um die Sprache zu ändern:\n"
        "1. Suchen Sie den Sprachschalter in der oberen Navigationsleiste.\n"
        "2. Klicken Sie darauf, um das Dropdown-Menü zu öffnen.\n"
        "3. Wählen Sie Ihre bevorzugte Sprache.\n"
        "4. Die gesamte Benutzeroberfläche wechselt sofort zur gewählten Sprache."
    )

    doc.add_heading("2.5  Dunkelmodus / Hellmodus", level=2)
    doc.add_paragraph(
        "Sie können zwischen einem hellen und dunklen Design wechseln:\n"
        "1. Finden Sie das Sonnen-/Mond-Symbol in der oberen Navigationsleiste.\n"
        "2. Klicken Sie darauf, um zwischen Hellmodus (weißer Hintergrund) und Dunkelmodus (dunkler Hintergrund) zu wechseln.\n"
        "3. Ihre Einstellung wird automatisch gespeichert."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # KAPITEL 3: LEITFADEN FÜR AUFTRAGGEBER
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("3  Leitfaden für Product Owner (Auftraggeber)", level=1)

    doc.add_heading("3.1  Meine Projekte — Übersicht", level=2)
    doc.add_paragraph(
        "Nach der Anmeldung als Product Owner sehen Sie die Seite 'Meine Projekte'. "
        "Hier werden alle Projekte aufgelistet, zu denen Sie eingeladen oder zugewiesen wurden."
    )
    doc.add_paragraph(
        "Jede Projektkarte zeigt:\n"
        "• Den Projektnamen\n"
        "• Den aktuellen Status (z. B. 'In Bearbeitung', 'Eingereicht')\n"
        "• Einen Fortschrittsbalken, der zeigt, wie viele Fragen beantwortet wurden\n"
        "• Eine Schaltfläche zum Fortfahren oder Überprüfen der Einreichung"
    )

    doc.add_heading("3.2  Anforderungsformular ausfüllen", level=2)
    doc.add_paragraph(
        "Klicken Sie auf ein Projekt, um das Anforderungsformular zu öffnen. "
        "Das Formular ist in logische Abschnitte unterteilt (z. B. Allgemeine Idee, "
        "Zielgruppe, Technische Funktionen, Budget)."
    )

    doc.add_heading("3.2.1  Schritt-für-Schritt-Formularnavigation", level=3)
    doc.add_paragraph(
        "1. Das Formular führt Sie Abschnitt für Abschnitt.\n"
        "2. Füllen Sie jedes Feld aus — Text, Auswahl, Datum oder Datei-Upload.\n"
        "3. Klicken Sie auf 'Weiter', um zum nächsten Abschnitt zu gelangen, oder 'Zurück'.\n"
        "4. Ihre Antworten werden automatisch gespeichert, während Sie tippen.\n"
        "5. Sie können das Formular jederzeit verlassen und später zurückkehren — Ihr Fortschritt bleibt erhalten."
    )

    doc.add_heading("3.2.2  Spracheingabe verwenden", level=3)
    doc.add_paragraph("Statt zu tippen, können Sie Ihre Antworten sprechen:")
    doc.add_paragraph(
        "1. Suchen Sie das Mikrofon-Symbol neben den Textfeldern.\n"
        "2. Klicken Sie auf das Mikrofon-Symbol.\n"
        "3. Erlauben Sie Ihrem Browser den Zugriff auf Ihr Mikrofon (nur beim ersten Mal).\n"
        "4. Sprechen Sie deutlich — Ihre Worte erscheinen in Echtzeit als Text.\n"
        "5. Klicken Sie auf die Stopp-Taste, wenn Sie fertig sind.\n"
        "6. Das System reinigt und verbessert Ihren gesprochenen Text automatisch."
    )
    doc.add_paragraph(
        "Tipp: Die Spracheingabe funktioniert am besten in einer ruhigen Umgebung. "
        "Sprechen Sie in vollständigen Sätzen für die besten Ergebnisse."
    )

    doc.add_heading("3.2.3  KI-Textverbesserung", level=3)
    doc.add_paragraph(
        "Das Portal enthält eine intelligente Textverbesserungsfunktion, unterstützt durch KI (Google Gemini). "
        "Wenn Sie Text eingeben (per Tastatur oder Sprache), kann das System:"
    )
    doc.add_paragraph("Grammatik- und Rechtschreibfehler korrigieren", style="List Bullet")
    doc.add_paragraph("Den professionellen Ton Ihres Textes verbessern", style="List Bullet")
    doc.add_paragraph("Ihre Gedanken in klare Anforderungsaussagen strukturieren", style="List Bullet")
    doc.add_paragraph(
        "Die Verbesserung erfolgt automatisch oder kann manuell ausgelöst werden. "
        "Sie sehen immer sowohl Ihren Originaltext als auch die verbesserte Version, "
        "sodass Sie die Änderungen annehmen oder verwerfen können."
    )

    doc.add_heading("3.2.4  KI-Nachfragen", level=3)
    doc.add_paragraph(
        "Nachdem Sie eine Frage beantwortet haben, kann die KI Nachfragen vorschlagen, "
        "um Ihnen zu helfen, mehr Details anzugeben. Diese erscheinen als hervorgehobene Karten "
        "unterhalb Ihrer Antwort."
    )
    doc.add_paragraph(
        "Sie können diese Nachfragen beantworten oder überspringen — sie sind optional, "
        "helfen aber dabei, bessere und vollständigere Anforderungen zu erstellen."
    )

    doc.add_heading("3.2.5  Dateien hochladen", level=3)
    doc.add_paragraph(
        "Bei einigen Fragen können oder müssen Dateien hochgeladen werden "
        "(z. B. Mockups, Wireframes oder Referenzdokumente).\n"
        "1. Klicken Sie auf den Upload-Bereich oder ziehen Sie eine Datei per Drag-and-Drop.\n"
        "2. Unterstützte Formate: PDF, Bilder (PNG, JPG) und gängige Office-Dokumente.\n"
        "3. Maximale Dateigröße: 50 MB pro Datei.\n"
        "4. Hochgeladene Dateien erscheinen als Miniaturvorschau unter der Frage."
    )

    doc.add_heading("3.3  KI-Interviewer verwenden", level=2)
    doc.add_paragraph(
        "Der KI-Interviewer ist eine alternative Möglichkeit, Ihre Anforderungen in einem "
        "natürlichen Gespräch zu formulieren, statt Formularfelder auszufüllen."
    )
    doc.add_paragraph(
        "1. Öffnen Sie Ihr Projektformular.\n"
        "2. Suchen Sie die Schaltfläche 'KI-Interview' oder 'Chat-Modus'.\n"
        "3. Wechseln Sie in den Interview-Modus.\n"
        "4. Ein freundlicher KI-Bot stellt Ihnen nacheinander Fragen.\n"
        "5. Sie können Ihre Antworten tippen oder per Spracheingabe geben.\n"
        "6. Die KI versteht den Kontext und stellt relevante Nachfragen.\n"
        "7. Nach Abschluss des Interviews werden alle Ihre Antworten automatisch "
        "den Formularfeldern zugeordnet."
    )
    doc.add_paragraph(
        "Tipp: Der KI-Interviewer ist ideal, wenn Sie einen dialogorientierten Ansatz bevorzugen "
        "oder Schwierigkeiten beim Ausfüllen strukturierter Formulare haben."
    )

    doc.add_heading("3.4  Überprüfen und Einreichen", level=2)
    doc.add_paragraph(
        "Vor der endgültigen Einreichung:\n"
        "1. Navigieren Sie zum Abschnitt 'Überprüfung' im Formular.\n"
        "2. Sie sehen eine vollständige Zusammenfassung aller Antworten, Abschnitt für Abschnitt.\n"
        "3. Prüfen Sie alles sorgfältig — nach der Einreichung können Antworten nicht einfach geändert werden.\n"
        "4. Wenn Sie etwas ändern möchten, klicken Sie auf den Abschnittsnamen, um zurückzugehen.\n"
        "5. Wenn Sie zufrieden sind, klicken Sie auf 'Einreichen'.\n"
        "6. Eine Bestätigungsmeldung erscheint. Ihre Einreichung ist nun für das WMC-Team sichtbar."
    )

    doc.add_heading("3.5  Nach der Einreichung", level=2)
    doc.add_paragraph(
        "Nach der Einreichung ändert sich Ihr Status auf 'Eingereicht'. "
        "Das WMC-Team prüft Ihre Anforderungen und kann:\n"
        "• Die Antwort als 'Geprüft' markieren (keine weitere Aktion Ihrerseits erforderlich)\n"
        "• Zusätzliche Informationen anfordern (Sie werden benachrichtigt)\n\n"
        "Sie können den Status Ihrer Einreichung jederzeit im Dashboard 'Meine Projekte' einsehen."
    )

    create_styled_table(doc,
        ["Status", "Bedeutung"],
        [
            ["Entwurf", "Sie haben das Formular begonnen, aber noch nicht eingereicht"],
            ["In Bearbeitung", "Sie arbeiten aktiv am Formular"],
            ["Eingereicht", "Sie haben eingereicht; das WMC-Team prüft"],
            ["Geprüft", "Das WMC-Team hat Ihre Einreichung geprüft"],
        ],
        tn(), "Status-Werte der Einreichungen"
    )

    doc.add_heading("3.6  Kontoeinstellungen", level=2)
    doc.add_paragraph(
        "Greifen Sie auf Ihre Kontoeinstellungen zu, indem Sie auf Ihr Profilsymbol in der "
        "oberen rechten Ecke klicken oder zur Seite 'Konto' navigieren."
    )
    doc.add_paragraph(
        "Hier können Sie:\n"
        "• Ihre zugewiesene Rolle und Kontoinformationen einsehen\n"
        "• Ihr Passwort ändern\n"
        "• Sehen, wie lange Sie bereits Mitglied sind"
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # KAPITEL 4: ADMIN-LEITFADEN
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("4  Leitfaden für Administratoren (WMC-Mitarbeiter)", level=1)

    doc.add_heading("4.1  Admin-Dashboard — Übersicht", level=2)
    doc.add_paragraph(
        "Nach der Anmeldung als Mitarbeiter oder Super-Admin landen Sie auf dem Admin-Dashboard. "
        "Es bietet einen schnellen Überblick über die Aktivitäten Ihrer Organisation."
    )
    doc.add_paragraph("Das Dashboard zeigt folgende Kennzahlen:")
    doc.add_paragraph("Gesamte Projekte — Anzahl aller Projekte in Ihrer Organisation", style="List Bullet")
    doc.add_paragraph("Aktive Projekte — derzeit in Bearbeitung befindliche Projekte", style="List Bullet")
    doc.add_paragraph("Gesamte Einreichungen — alle eingegangenen Antworten von Auftraggebern", style="List Bullet")
    doc.add_paragraph("Ausstehende Prüfungen — Einreichungen, die auf Ihre Prüfung warten (rot hervorgehoben)", style="List Bullet")

    create_styled_table(doc,
        ["Menüpunkt", "Beschreibung"],
        [
            ["Dashboard", "Übersicht mit Kennzahlen und letzter Aktivität"],
            ["Projekte", "Anforderungsprojekte erstellen, verwalten und zuweisen"],
            ["Vorlagen", "Fragevorlagen für Projektformulare gestalten"],
            ["Einreichungen", "Kundeneinreichungen prüfen, filtern und exportieren"],
            ["Archiv", "Abgeschlossene oder inaktive Projekte einsehen"],
            ["Einstellungen", "Organisations- und Systemeinstellungen"],
        ],
        tn(), "Navigationsübersicht — Administrator"
    )

    doc.add_heading("4.2  Projekte verwalten", level=2)

    doc.add_heading("4.2.1  Neues Projekt erstellen", level=3)
    doc.add_paragraph(
        "1. Gehen Sie zu 'Projekte' im Seitenmenü.\n"
        "2. Klicken Sie auf 'Neues Projekt'.\n"
        "3. Füllen Sie die Projektdetails aus:\n"
        "   • Projektname: Ein klarer, beschreibender Name (z. B. 'Reise-App Anforderungen')\n"
        "   • Slug: Ein automatisch generierter URL-freundlicher Bezeichner\n"
        "   • Frist: Anzahl der Tage bis zur Fertigstellung\n"
        "   • Vorlage: Wählen Sie eine Fragenvorlage\n"
        "   • Begrüßungstext: Eine Willkommensnachricht für den Auftraggeber\n"
        "4. Klicken Sie auf 'Projekt erstellen'.\n"
        "5. Das Projekt ist nun bereit für Einreichungen."
    )

    doc.add_heading("4.2.2  Teammitglieder einladen", level=3)
    doc.add_paragraph(
        "Nach dem Erstellen eines Projekts können Sie Auftraggeber einladen:\n"
        "1. Öffnen Sie die Projektdetailseite.\n"
        "2. Klicken Sie auf 'Mitglied einladen'.\n"
        "3. Geben Sie die E-Mail-Adresse der Person ein.\n"
        "4. Wählen Sie die Rolle (Product Owner oder Client).\n"
        "5. Eine Magic-Link-Einladung wird automatisch per E-Mail versendet.\n"
        "6. Die Einladung ist 7 Tage gültig."
    )

    doc.add_heading("4.2.3  Projektvorschläge genehmigen", level=3)
    doc.add_paragraph(
        "Wenn Auftraggeber Projektideen einreichen, die genehmigt werden müssen:\n"
        "1. Öffnen Sie die Projektdetailseite.\n"
        "2. Prüfen Sie Beschreibung und Umfang des Projekts.\n"
        "3. Klicken Sie auf 'Genehmigen', um das Projekt zu aktivieren, oder\n"
        "4. Klicken Sie auf 'Änderungen anfordern' mit einer Anmerkung."
    )

    doc.add_heading("4.3  Vorlagen verwalten", level=2)
    doc.add_paragraph(
        "Vorlagen definieren die Struktur des Anforderungsformulars. "
        "Eine Vorlage besteht aus Abschnitten, und jeder Abschnitt enthält Fragen."
    )

    doc.add_heading("4.3.1  Vorlage erstellen", level=3)
    doc.add_paragraph(
        "1. Gehen Sie zu 'Vorlagen' im Seitenmenü.\n"
        "2. Klicken Sie auf 'Neue Vorlage'.\n"
        "3. Vergeben Sie einen Namen und eine optionale Beschreibung.\n"
        "4. Fügen Sie Abschnitte hinzu (z. B. 'Allgemeine Informationen', 'Technische Anforderungen').\n"
        "5. Fügen Sie innerhalb jedes Abschnitts Fragen verschiedener Typen hinzu."
    )

    create_styled_table(doc,
        ["Fragetyp", "Beschreibung", "Beispielverwendung"],
        [
            ["Kurztext", "Einzeilige Texteingabe", "Projektname, Firmenname"],
            ["Langtext", "Mehrzeiliges Textfeld", "Detaillierte Beschreibungen, Ideen"],
            ["Einfachauswahl", "Optionsschaltflächen — eine Option wählen", "Priorität (Hoch/Mittel/Niedrig)"],
            ["Mehrfachauswahl", "Kontrollkästchen — mehrere Optionen wählen", "Gewünschte Plattformen (Web, iOS, Android)"],
            ["Dropdown", "Aus einer Liste auswählen", "Branche"],
            ["Datum", "Datumsauswahl", "Gewünschter Starttermin"],
            ["Datei-Upload", "Dokumente oder Bilder anhängen", "Mockups, Wireframes, Referenz-PDFs"],
        ],
        tn(), "Fragetypen in Vorlagen"
    )

    doc.add_heading("4.3.2  Abschnitte und Fragen umsortieren", level=3)
    doc.add_paragraph(
        "Sie können die Reihenfolge per Drag-and-Drop ändern:\n"
        "1. Bewegen Sie die Maus über den Ziehgriff (≡-Symbol).\n"
        "2. Klicken und halten Sie, dann ziehen Sie an die gewünschte Position.\n"
        "3. Lassen Sie los, um abzulegen.\n"
        "4. Änderungen werden automatisch gespeichert."
    )

    doc.add_heading("4.4  Einreichungen prüfen", level=2)
    doc.add_paragraph(
        "Wenn Auftraggeber ihre Anforderungen einreichen, können Sie diese im Bereich "
        "'Einreichungen' prüfen."
    )
    doc.add_paragraph(
        "1. Gehen Sie zu 'Einreichungen' im Seitenmenü.\n"
        "2. Sie sehen eine Liste aller Einreichungen mit Status und Fortschritt.\n"
        "3. Nutzen Sie die Filter:\n"
        "   • Nach Status filtern (Entwurf, In Bearbeitung, Eingereicht, Geprüft)\n"
        "   • Nach Name oder Projektname suchen\n"
        "4. Klicken Sie auf eine Einreichung für die vollständigen Details.\n"
        "5. Prüfen Sie jede Antwort sorgfältig.\n"
        "6. Markieren Sie die Einreichung als 'Geprüft', wenn Sie fertig sind."
    )

    doc.add_heading("4.4.1  Einreichungen exportieren", level=3)
    doc.add_paragraph(
        "Sie können Daten für die Offline-Verarbeitung exportieren:\n"
        "1. Klicken Sie in der Liste auf 'CSV Exportieren'.\n"
        "2. Eine CSV-Datei wird heruntergeladen.\n"
        "3. Öffnen Sie sie in Microsoft Excel, Google Sheets oder einer anderen Tabellenkalkulation."
    )

    doc.add_heading("4.5  Archiv", level=2)
    doc.add_paragraph(
        "Im Archiv werden abgeschlossene oder deaktivierte Projekte gespeichert. "
        "Sie können weiterhin alle Details und Einreichungen archivierter Projekte einsehen."
    )

    doc.add_heading("4.6  Einstellungen", level=2)
    doc.add_paragraph(
        "Die Einstellungsseite ermöglicht Administratoren die Konfiguration von "
        "organisationsweiten Einstellungen wie Organisationsname, Standardsprache "
        "und Benachrichtigungspräferenzen."
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # KAPITEL 5: NAVIGATIONSREFERENZ
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("5  Navigationsreferenz", level=1)

    doc.add_heading("5.1  Auftraggeber-Navigation", level=2)
    create_styled_table(doc,
        ["Seite", "Zugang", "Was Sie tun können"],
        [
            ["Meine Projekte", "Automatisch nach Anmeldung", "Alle zugewiesenen Projekte einsehen"],
            ["Formular (Ausfüll-Modus)", "Projekt anklicken", "Fragen Abschnitt für Abschnitt beantworten"],
            ["Formular (Interview-Modus)", "Modus im Formular wechseln", "Fragen im KI-Gespräch beantworten"],
            ["Überprüfung", "Letzter Schritt im Formular", "Alle Antworten vor dem Einreichen prüfen"],
            ["Konto", "Profilsymbol klicken", "Kontoinformationen einsehen, Passwort ändern"],
        ],
        tn(), "Navigationsübersicht — Auftraggeber"
    )

    doc.add_heading("5.2  Tastenkombinationen", level=2)
    create_styled_table(doc,
        ["Tastenkombination", "Aktion"],
        [
            ["Tab", "Zum nächsten Formularfeld springen"],
            ["Umschalt + Tab", "Zum vorherigen Formularfeld springen"],
            ["Eingabe", "Auswahl bestätigen oder aktuelles Feld absenden"],
            ["Strg + S", "Aktuellen Fortschritt speichern (in unterstützten Browsern)"],
            ["Esc", "Dialoge oder Pop-ups schließen"],
        ],
        tn(), "Tastenkombinationen"
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # KAPITEL 6: UNTERSTÜTZTE SPRACHEN
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("6  Unterstützte Sprachen", level=1)
    doc.add_paragraph(
        "Das WMC-Anforderungsportal unterstützt 25 Sprachen. Die Standardsprache ist Deutsch. "
        "Alle Oberflächenelemente, Formularbeschriftungen und KI-generierten Texte erscheinen "
        "in der gewählten Sprache."
    )
    create_styled_table(doc,
        ["Code", "Sprache", "Code", "Sprache"],
        [
            ["de", "Deutsch", "nl", "Niederländisch"],
            ["en", "Englisch", "no", "Norwegisch"],
            ["bg", "Bulgarisch", "pl", "Polnisch"],
            ["cs", "Tschechisch", "pt", "Portugiesisch"],
            ["da", "Dänisch", "ro", "Rumänisch"],
            ["el", "Griechisch", "ru", "Russisch"],
            ["es", "Spanisch", "sk", "Slowakisch"],
            ["et", "Estnisch", "sl", "Slowenisch"],
            ["fi", "Finnisch", "sv", "Schwedisch"],
            ["fr", "Französisch", "tr", "Türkisch"],
            ["hr", "Kroatisch", "hu", "Ungarisch"],
            ["it", "Italienisch", "lt", "Litauisch"],
            ["lv", "Lettisch", "", ""],
        ],
        tn(), "Unterstützte Sprachen"
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # KAPITEL 7: DATENSCHUTZ UND SICHERHEIT
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("7  Datenschutz und Sicherheit", level=1)

    doc.add_heading("7.1  DSGVO-Konformität", level=2)
    doc.add_paragraph(
        "Das WMC-Anforderungsportal ist vollständig konform mit der Europäischen "
        "Datenschutz-Grundverordnung (DSGVO). Wichtige Maßnahmen umfassen:"
    )
    doc.add_paragraph("Alle Daten werden auf Servern innerhalb der Europäischen Union gespeichert", style="List Bullet")
    doc.add_paragraph("Die Datenübertragung ist durch HTTPS/TLS verschlüsselt", style="List Bullet")
    doc.add_paragraph("Der Zugriff wird durch rollenbasierte Berechtigungen kontrolliert", style="List Bullet")
    doc.add_paragraph("Benutzer können jederzeit die Löschung ihrer Daten beantragen", style="List Bullet")
    doc.add_paragraph("KI-Verarbeitung erfolgt sicher — keine Daten werden für Modelltraining verwendet", style="List Bullet")

    doc.add_heading("7.2  Datenspeicherung", level=2)
    doc.add_paragraph(
        "Das Portal verwendet Supabase als Backend-Datenbankdienst. "
        "Supabase bietet PostgreSQL-Datenbanken auf Enterprise-Niveau mit integrierter "
        "Authentifizierung, zeilenbasierter Sicherheit und automatischen Backups."
    )

    doc.add_heading("7.3  Authentifizierungssicherheit", level=2)
    doc.add_paragraph(
        "Die Anmeldung wird geschützt durch:\n"
        "• Mehrere Optionen (Magic Links, Passwörter)\n"
        "• Verschlüsselte Passwortspeicherung (bcrypt-Hashing)\n"
        "• Token-basierte Sitzungsverwaltung (JWT)\n"
        "• Automatischer Sitzungsablauf nach Inaktivität"
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # KAPITEL 8: FEHLERBEHEBUNG
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("8  Fehlerbehebung", level=1)
    doc.add_paragraph(
        "Wenn Sie Probleme bei der Nutzung des Portals haben, versuchen Sie die folgenden "
        "Lösungen, bevor Sie den Support kontaktieren."
    )

    create_styled_table(doc,
        ["Problem", "Mögliche Ursache", "Lösung"],
        [
            [
                "Anmeldung nicht möglich",
                "Falsche E-Mail oder Passwort",
                "Überprüfen Sie Ihre Zugangsdaten. Nutzen Sie 'Passwort vergessen' zum Zurücksetzen. "
                "Bei Magic Link: Prüfen Sie den Spam-Ordner."
            ],
            [
                "Seite lädt nicht",
                "Internetproblem oder Browser-Cache",
                "Prüfen Sie Ihre Internetverbindung. Versuchen Sie, den Browser-Cache zu löschen "
                "(Strg+Umschalt+Entf) oder nutzen Sie ein privates Fenster."
            ],
            [
                "Spracheingabe funktioniert nicht",
                "Mikrofon nicht erlaubt",
                "Ihr Browser benötigt die Berechtigung für das Mikrofon. "
                "Klicken Sie auf das Schloss-Symbol in der Adressleiste und erlauben Sie den Zugriff."
            ],
            [
                "KI-Textverbesserung funktioniert nicht",
                "Text zu kurz oder unklar",
                "Geben Sie mindestens 2-3 vollständige Sätze ein. Die KI funktioniert am besten "
                "mit längeren, detaillierteren Eingaben."
            ],
            [
                "Datei-Upload fehlgeschlagen",
                "Datei zu groß oder Format nicht unterstützt",
                "Maximale Dateigröße: 50 MB. Unterstützte Formate: PDF, PNG, JPG "
                "und gängige Office-Dokumente."
            ],
            [
                "Fortschritt wurde nicht gespeichert",
                "Sitzung abgelaufen",
                "Das System speichert automatisch, aber bei längerer Inaktivität werden Sie "
                "abgemeldet. Melden Sie sich erneut an — Ihr Fortschritt sollte weitgehend erhalten sein."
            ],
            [
                "Sprache hat sich unerwartet geändert",
                "Browser-Spracherkennung",
                "Das Portal erkennt automatisch Ihre Browser-Sprache. Nutzen Sie den "
                "Sprachschalter in der oberen Leiste zur manuellen Auswahl."
            ],
            [
                "Mein Projekt wird nicht angezeigt",
                "Nicht dem Projekt zugewiesen",
                "Kontaktieren Sie Ihren WMC-Administrator, um zu bestätigen, dass Sie "
                "zum richtigen Projekt eingeladen wurden."
            ],
        ],
        tn(), "Fehlerbehebung"
    )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # KAPITEL 9: FAQ
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("9  Häufig gestellte Fragen (FAQ)", level=1)

    faqs = [
        ("Kann ich meine Antworten nach dem Einreichen bearbeiten?",
         "Nach der Einreichung sind Ihre Antworten für die Prüfung durch das WMC-Team gesperrt. "
         "Wenn Sie Änderungen vornehmen müssen, kontaktieren Sie Ihren WMC-Projektmanager, "
         "der das Formular für Sie wieder öffnen kann."),
        ("Sind meine Daten sicher?",
         "Ja. Alle Daten werden bei der Übertragung und Speicherung verschlüsselt. "
         "Das Portal ist DSGVO-konform und Ihre Daten werden auf EU-Servern gespeichert. "
         "KI-Funktionen verwenden Ihre Daten nicht zu Trainingszwecken."),
        ("Kann ich das Portal auf meinem Smartphone nutzen?",
         "Ja. Das Portal ist responsive und funktioniert auf Smartphones und Tablets. "
         "Für die beste Erfahrung empfehlen wir einen Desktop- oder Laptop-Computer, "
         "insbesondere bei längeren Formularen."),
        ("Was passiert bei einem Internetausfall während des Ausfüllens?",
         "Ihre Antworten werden automatisch gespeichert. Wenn die Verbindung abbricht, "
         "stellen Sie sie wieder her und laden Sie die Seite neu — Ihr Fortschritt sollte erhalten sein."),
        ("Muss ich Software installieren?",
         "Nein. Das Portal läuft vollständig in Ihrem Webbrowser. "
         "Keine Downloads oder Installationen erforderlich."),
        ("Wie viele Sprachen werden unterstützt?",
         "Das Portal unterstützt 25 Sprachen, darunter Deutsch, Englisch, Französisch, Spanisch, "
         "Italienisch, Polnisch, Türkisch und viele weitere. Siehe Kapitel 6 für die vollständige Liste."),
        ("Wer kann meine eingereichten Anforderungen sehen?",
         "Nur die WMC-Mitarbeiter und Administratoren, die Ihrem Projekt zugewiesen sind, "
         "können Ihre Einreichungen einsehen. Andere Auftraggeber können Ihre Daten nicht sehen."),
        ("Funktioniert die Spracheingabe in jeder Sprache?",
         "Die Spracheingabe hängt von den Spracherkennungsfähigkeiten Ihres Browsers ab. "
         "Die meisten modernen Browser unterstützen die wichtigsten europäischen Sprachen. "
         "Die besten Ergebnisse erzielen Sie mit Chrome."),
    ]
    for q, a in faqs:
        p = doc.add_paragraph()
        run = p.add_run(f"F: {q}")
        run.bold = True
        doc.add_paragraph(f"A: {a}")
        doc.add_paragraph("")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # KAPITEL 10: GLOSSAR
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("10  Glossar", level=1)
    terms = [
        ("Anforderungsportal", "Die Webanwendung zur strukturierten Erfassung von Projektanforderungen."),
        ("KI-Interviewer", "Ein chatbasierter Assistent, der Ihnen Fragen stellt, um Anforderungen zu sammeln."),
        ("KI-Textverbesserung", "Eine automatische Funktion, die Grammatik und Ton Ihres Textes verbessert."),
        ("Dashboard", "Die Übersichtsseite mit Kennzahlen und aktueller Aktivität."),
        ("DSGVO", "Datenschutz-Grundverordnung — europäische Verordnung zum Schutz personenbezogener Daten."),
        ("Magic Link", "Ein sicherer, einmaliger Anmeldelink, der an Ihre E-Mail gesendet wird."),
        ("Product Owner", "Die Person, die die Produktidee hat und Anforderungen liefert."),
        ("Einreichung", "Die Gesamtheit der Antworten, die ein Auftraggeber für ein Projekt einreicht."),
        ("Abschnitt", "Eine Gruppe zusammenhängender Fragen innerhalb einer Vorlage."),
        ("Slug", "Eine URL-freundliche Version eines Projektnamens (z. B. 'reise-app-anforderungen')."),
        ("Mitarbeiter", "WMC-Berater, die Projekte verwalten und Einreichungen prüfen."),
        ("Super-Admin", "Ein Systemadministrator mit Vollzugriff auf alle Funktionen."),
        ("Vorlage", "Eine vordefinierte Struktur aus Abschnitten und Fragen für die Anforderungserfassung."),
        ("Spracheingabe", "Sprechen statt Tippen — das Portal wandelt Ihre Sprache in Text um."),
    ]
    for term, defn in terms:
        p = doc.add_paragraph()
        run = p.add_run(f"{term}: ")
        run.bold = True
        p.add_run(defn)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # KAPITEL 11: SUPPORT UND KONTAKT
    # ════════════════════════════════════════════════════════════════════════
    doc.add_heading("11  Support und Kontakt", level=1)
    doc.add_paragraph("Bei weiteren Fragen kontaktieren Sie bitte das WMC-Support-Team:")
    doc.add_paragraph("E-Mail: support@wamocon.com", style="List Bullet")
    doc.add_paragraph("Website: https://wamocon.com", style="List Bullet")
    doc.add_paragraph("Telefon: Kontaktieren Sie Ihren Projektmanager für direkte Unterstützung", style="List Bullet")
    doc.add_paragraph("")
    doc.add_paragraph(
        "Bitte geben Sie bei der Kontaktaufnahme an:\n"
        "• Ihre E-Mail-Adresse (für die Anmeldung verwendet)\n"
        "• Den Projektnamen\n"
        "• Eine Beschreibung des Problems\n"
        "• Screenshots, falls verfügbar"
    )
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("— Ende des Benutzerhandbuchs —")
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.size = Pt(10)


# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────

def main():
    # English
    doc_en = Document()
    build_english(doc_en)
    en_path = os.path.join(OUTPUT_DIR, "WMC_Anforderungsportal_User_Manual_EN.docx")
    doc_en.save(en_path)
    print(f"[OK] English manual saved: {en_path}")

    # German
    doc_de = Document()
    build_german(doc_de)
    de_path = os.path.join(OUTPUT_DIR, "WMC_Anforderungsportal_Benutzerhandbuch_DE.docx")
    doc_de.save(de_path)
    print(f"[OK] German manual saved: {de_path}")

    print(f"\nBoth documents are in: {OUTPUT_DIR}")
    print("Tip: Open in Word and press Ctrl+A > F9 to update the Table of Contents.")


if __name__ == "__main__":
    main()
